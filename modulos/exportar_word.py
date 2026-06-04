from io import BytesIO
from docx import Document


def generar_word(estudiante, informe):
    documento = Document()

    documento.add_heading("Informe Comparativo", level=1)

    documento.add_paragraph(f"Estudiante: {estudiante}")

    documento.add_heading("Resultados", level=2)

    documento.add_paragraph(str(informe))

    archivo = BytesIO()

    documento.save(archivo)

    archivo.seek(0)

    return archivo